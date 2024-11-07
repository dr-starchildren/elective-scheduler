from itertools import combinations
import random
import csv
import pandas as pd 
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches

def get_spreadsheet_data(filename='scheduling.xlsx', sheet_name=''):
    df = pd.read_excel(io=filename, sheet_name=sheet_name)
    #data = [row.dropna().tolist() for _, row in df.iterrows()]
    return df

def get_rankings(filename):
    df = pd.read_csv(filename)
    rankings = [row.tolist() for _, row in df.iterrows()]
    return rankings

def generate_random_ranking(electives, all):
    num_electives = len(electives)
    
    # Decide how many activities to rank (at least 1, at most all)
    if all:
        num_ranked = num_electives
    else:
        num_ranked = random.randint(12, num_electives)
    # Randomly select a subset of activities to rank
    ranked_electives = random.sample(electives, num_ranked)
    
    # Assign ranks from 1 to num_ranked
    rankings = {}
    available_ranks = list(range(1, num_ranked + 1))
    random.shuffle(available_ranks)
    
    for elective in ranked_electives:
        rankings[elective] = available_ranks.pop()
    
    # Assign None to activities that are not ranked
    for elective in electives:
        if elective not in rankings:
            rankings[elective] = None
    
    return rankings

def create_ranking_csv(output, names, all=False):
    electives = get_spreadsheet_data()
    elective_names = [elective[0] for elective in electives]
    camper_rankings = {}
    for name in names:
        rankings = generate_random_ranking(elective_names, all)
        camper_rankings[name] = sorted(rankings.items(), key=lambda item: item[0])

    with open(output, 'w', newline='') as w:
        w.flush()
        writer = csv.writer(w,delimiter=',')
        writer.writerow(['Name']+[ elective[0] for elective in list(camper_rankings.values())[0]])
        for name in camper_rankings.keys():
            writer.writerow([name]+[ elective[1] for elective in camper_rankings[name]])


"""
    Returns the abbreviated name of the day of the week based on an integer input.

    This function maps an integer representing the day of the week to its corresponding
    abbreviated weekday name. If the input integer is out of the valid range (0–6),
    it returns 'Unknown'.
"""
def day_of_week(day: int) -> str:
    days = {0: 'Mon', 1: 'Tues', 2: 'Wed', 3: 'Thur', 4: 'Fri', 5: 'Sat', 6: 'Sun'}
    return days.get(day, 'Unknown')
    
# In case we need to generate random data 
#create_ranking_csv('rankings.csv', campers)
#create_ranking_csv('counselor_rankings.csv', counselors, True)


"""
    Returns a formatted string representation of a camper's name for printing.

    The function retrieves a camper's full name from the `all_campers` dictionary using
    the provided `camper_id`, and formats the name according to these rules:
    - If the name has a first name and last name, return the first name and the first initial of the last name.
    - If the name has two first names and a last name, return the first and second names along with the first initial of the last name.
    - If the name only has a first name, return the first name.

    Parameters:
    camper_id (int or str): The ID of the camper, used as a key in the `all_campers` dictionary.

    Returns:
    str: A formatted string with the camper's name, tailored for readability.
"""
def printable_camper(camper_id, all_campers):
    camper = all_campers[camper_id]
    split = camper.split(' ')
    # If you have two first names take both and first initial
    if len(split) > 2:
        return f"{split[0]} {split[1]} {split[2][0]}"
    # Take first name and first initial
    if len(split) > 1:
        return f"{split[0]} {split[1][0]}"
    # If you don't have a last name, just take the first name
    return split[0]

def resize_columns(sheet):
    for col in sheet.columns:
        max_length = 0
        column = col[0].column_letter 
        for cell in col:
            try:
                max_length = max(max_length, len(str(cell.value)))
            except:
                pass
        adjusted_width = max_length + 2  # Add extra space for padding
        sheet.column_dimensions[column].width = adjusted_width


"""
    Adds formatted content for an elective's schedule to a Word document.

    This function generates a formatted section in a Word document for a specific elective's
    schedule, including the elective name, scheduled day, counselor names, and a list of campers
    with checkboxes. Each section is separated by a page break for clarity.

    Function Details:
    - Creates a bold title with the elective name and the day of the week.
    - Adds counselor names beneath the elective name with a centered alignment.
    - Lists each camper’s name in alphabetical order by last name, prefixed by a checkbox symbol ("☐").
    - Adds a page break after each elective section to separate it from subsequent content.
"""
def word_doc_output(doc, elective, campers, day, counselors):
    
    p = doc.add_paragraph()
    run = p.add_run(f"{elective} ({day_of_week(day)})")
    run.bold = True
    font = run.font
    font.size = Pt(20)
    run.add_break()
    run = p.add_run(", ".join(counselors))
    font = run.font
    font.size = Pt(16)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add campers with checkboxes
    camper_names_sorted = sorted(campers, key=lambda name: name.split(' ')[-1])
    for camper in camper_names_sorted:
        p = doc.add_paragraph()
        run = p.add_run()
        font = run.font
        font.size = Pt(16)
        run.add_text('☐  ' + camper)
    
    doc.add_page_break()
